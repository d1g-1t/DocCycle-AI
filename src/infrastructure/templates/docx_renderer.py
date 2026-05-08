"""DOCX renderer — generates Word documents from template DSL + variables."""
from __future__ import annotations

import re
from io import BytesIO
from typing import Any

import structlog

log = structlog.get_logger(__name__)


class DocxRenderer:
    """Render a contract template into a DOCX file.

    Converts HTML or plain text output of template rendering
    into a formatted Word document suitable for signing.
    """

    def render(self, html: str, metadata: dict[str, Any] | None = None) -> bytes:
        """Convert rendered HTML to DOCX bytes."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError:
            log.warning("docx.not_available", hint="pip install python-docx")
            return b""

        doc = Document()

        # Optional header with metadata
        if metadata:
            title = metadata.get("title", "Contract")
            doc.add_heading(title, level=0)
            if metadata.get("contract_type"):
                doc.add_paragraph(
                    f"Type: {metadata['contract_type']}", style="Subtitle"
                )

        # Strip HTML tags and convert to paragraphs
        text = re.sub(r"<br\s*/?>", "\n", html)
        text = re.sub(r"<h[1-6][^>]*>(.*?)</h[1-6]>", r"\n\1\n", text)
        text = re.sub(r"<li[^>]*>(.*?)</li>", r"• \1\n", text)
        text = re.sub(r"<[^>]+>", "", text)

        for paragraph in text.split("\n"):
            stripped = paragraph.strip()
            if not stripped:
                continue

            # Detect headings (ALL CAPS or numbered sections)
            if stripped.isupper() and len(stripped) < 100:
                doc.add_heading(stripped, level=1)
            elif re.match(r"^\d+\.\s", stripped):
                p = doc.add_paragraph(stripped)
                p.style = "List Number"
            elif stripped.startswith("• "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

        buf = BytesIO()
        doc.save(buf)
        log.debug("docx_render.ok", size=buf.tell())
        return buf.getvalue()
