"""Jinja2 template rendering adapter."""
from __future__ import annotations

import hashlib
from typing import Any
from uuid import UUID

import structlog
from jinja2 import BaseLoader, Environment, TemplateSyntaxError, UndefinedError
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from src.application.interfaces.i_template_renderer import ITemplateRenderer, RenderedDocument
from src.infrastructure.database.models.contract_template_model import (
    ContractTemplateModel,
    TemplateVersionModel,
)

log = structlog.get_logger(__name__)

_jinja_env = Environment(
    loader=BaseLoader(),
    autoescape=True,
    trim_blocks=True,
    lstrip_blocks=True,
    keep_trailing_newline=False,
)


class Jinja2TemplateRenderer(ITemplateRenderer):
    """Renders a template version's body with Jinja2 + variables."""

    def __init__(self, session: AsyncSession) -> None:
        self._session = session

    async def render(
        self,
        template_id: UUID,
        version: int,
        variables: dict[str, object],
        *,
        include_docx: bool = False,
    ) -> RenderedDocument:
        """Render template by ID + version number; returns HTML and optional DOCX."""
        tv = await self._session.scalar(
            select(TemplateVersionModel).where(
                TemplateVersionModel.template_id == template_id,
                TemplateVersionModel.version_number == version,
            )
        )
        if tv is None:
            raise ValueError(f"TemplateVersion {template_id} v{version} not found")

        body = tv.body_dsl or ""
        try:
            tpl = _jinja_env.from_string(body)
            rendered_html = tpl.render(**variables)
        except (TemplateSyntaxError, UndefinedError) as exc:
            log.error("template_render.error", template_id=str(template_id), error=str(exc))
            raise ValueError(f"Template rendering failed: {exc}") from exc

        docx_bytes: bytes | None = None
        if include_docx:
            docx_bytes = self._generate_docx(rendered_html)

        log.debug("template_render.ok", template_id=str(template_id), length=len(rendered_html))
        return RenderedDocument(html=rendered_html, docx_bytes=docx_bytes)

    async def validate_variables(
        self,
        template_id: UUID,
        version: int,
        variables: dict[str, object],
    ) -> list[str]:
        """Return list of missing or invalid variable names."""
        tv = await self._session.scalar(
            select(TemplateVersionModel).where(
                TemplateVersionModel.template_id == template_id,
                TemplateVersionModel.version_number == version,
            )
        )
        if tv is None:
            return [f"TemplateVersion {template_id} v{version} not found"]

        required_vars: list[str] = [v.get("name", "") for v in (tv.variables or [])]
        missing = [v for v in required_vars if v and v not in variables]
        return missing

    @staticmethod
    def _generate_docx(html: str) -> bytes:
        """Generate a basic DOCX from HTML content."""
        try:
            from docx import Document
            from io import BytesIO

            doc = Document()
            # Strip HTML tags for plain-text DOCX fallback
            import re
            text = re.sub(r"<[^>]+>", "", html)
            for paragraph in text.split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            buf = BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except ImportError:
            log.warning("docx.not_installed", hint="pip install python-docx")
            return b""
